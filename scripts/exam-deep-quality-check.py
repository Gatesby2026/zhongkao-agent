#!/usr/bin/env python3
"""
北京中考数学试卷 — 深度质量检查
Layer A: 题目文本质量抽检（占位符/公式/空文本）
Layer B: 模块-题型一致性检查
Layer C: 答案提取与比对（解析版docx）

用法:
    python3 scripts/exam-deep-quality-check.py [--subject math] [--layers A B C]
"""

import os
import re
import sys
import yaml
import argparse
from pathlib import Path
from collections import defaultdict

# ============================================================
# 配置
# ============================================================

ROOT = Path(__file__).parent.parent
ORIG_DIR = ROOT / "knowledge-original"
KB_DIR = ROOT / "knowledge-base"

SUBJECTS = {
    "math":    ("北京中考数学模拟卷", "mock-exams/math/beijing"),
    "chinese": ("北京中考语文模拟卷", "mock-exams/chinese/beijing"),
    "english": ("北京中考英语模拟卷", "mock-exams/english/beijing"),
    "physics": ("北京中考物理模拟卷", "mock-exams/physics/beijing"),
    "politics":("北京中考道法模拟卷", "mock-exams/politics/beijing"),
}

DISTRICT_MAP = {
    "海淀":"haidian","西城":"xicheng","东城":"dongcheng","朝阳":"chaoyang",
    "丰台":"fengtai","石景山":"shijingshan","通州":"tongzhou","大兴":"daxing",
    "房山":"fangshan","顺义":"shunyi","门头沟":"mentougou","密云":"miyun",
    "平谷":"pinggu","燕山":"yanshan","延庆":"yanqing",
}
EXAM_TYPE_MAP = {"一模":"yi","二模":"er","三模":"san","真题":"zhenti","开卷":"kaikao"}


# ============================================================
# Layer A: 题目文本质量
# ============================================================

def layer_a_text_quality(subject):
    """检查题目文本中的占位符、空文本、LaTeX配对等问题"""
    _, yaml_subfolder = SUBJECTS[subject]
    yaml_path = KB_DIR / yaml_subfolder

    issues = []

    for yf in sorted(yaml_path.glob("*.yaml")):
        data = yaml.safe_load(yf.read_text(encoding="utf-8"))
        if not data:
            continue
        label = '%s %s %s' % (data.get("year",""), data.get("district",""), data.get("exam_type",""))

        for q in data.get("questions", []):
            qid = q.get("id", "?")
            qtype = q.get("type", "")
            text = q.get("question", "")

            # 1. 空文本
            if not text.strip():
                issues.append((label, qid, qtype, "空文本", text))
                continue

            # 2. 未解析的占位符（排除OCR粘连产物）
            # 粘连判断：[图]前方3字符内是数字/字母，说明是"3[图]"类OCR产物，非真实缺失
            for m in re.finditer(r"\[图\]", text):
                pos = m.start()
                before = text[max(0, pos-3):pos].strip()
                is_ocr_artifact = bool(re.match(r"\d$", before))  # 数字结尾如"3["
                if not is_ocr_artifact:
                    issues.append((label, qid, qtype, "未解析图片占位符", text[:80]))

            # 3. LaTeX 公式配对检查
            dollar_count = text.count("$")
            if dollar_count % 2 != 0:
                issues.append((label, qid, qtype, "LaTeX公式$奇数个（配对不完整）", text[:80]))

            # 4. 文本过短（选择题/填空题少于5字视为可疑）
            if qtype in ("选择", "填空", "单选", "多选") and len(text.strip()) < 5:
                issues.append((label, qid, qtype, "题目文本过短(<5字)", text[:80]))

    # 输出
    print(f"\n{'='*60}")
    print(f"Layer A 题目文本质量 — {subject}")
    print(f"{'='*60}")

    # 按问题类型分组统计
    by_type = defaultdict(list)
    for item in issues:
        by_type[item[3]].append(item)

    print(f"  共发现 {len(issues)} 处文本问题")
    for ptype, items in sorted(by_type.items(), key=lambda x: -len(x[1])):
        print(f"\n  [{ptype}] ({len(items)}处)")
        for label, qid, qtype, _, snippet in items[:5]:
            print(f"    {label} 第{qid}题({qtype}): {snippet[:60]}")
        if len(items) > 5:
            print(f"    ... 还有 {len(items)-5} 处")

    if not issues:
        print(f"  ✅ 无文本问题")


# ============================================================
# Layer B: 模块-题型一致性
# ============================================================

# 各科题型 → 预期出现的模块（非强制，只标记异常）
MATH_JIADA_PATTERNS = {
    "数与式": ["numbersAndExpressions"],
    "方程": ["equationsAndInequalities"],
    "函数": ["functions"],
    "三角": ["triangles"],
    "四边": ["quadrilaterals"],
    "圆": ["circles"],
    "几何": ["geometryComprehensive", "triangles", "quadrilaterals"],
    "统计": ["statisticsAndProbability"],
    "概率": ["statisticsAndProbability"],
}

# 北京中考数学各题号的标准题型范围
# 选择题(1-8): 数与式、方程、函数(反比例)、概率统计
# 填空题(9-16): 几何为主
# 解答题(17-20): 基础解答（方程/函数）
# 解答题(21-25): 中档（几何/函数）
# 解答题(26-28): 压轴（新定义/几何综合）

MATH_ANSWER_QIDS = set(range(17, 29))  # 解答题题号


def layer_b_module_type_consistency(subject):
    """检查模块标注与题型是否匹配"""
    if subject != "math":
        print(f"\n{'='*60}")
        print(f"Layer B 模块-题型一致性 — {subject} (仅支持数学)")
        print(f"  暂未实现")
        return

    yaml_path = KB_DIR / "mock-exams/math/beijing"
    issues = []

    for yf in sorted(yaml_path.glob("*.yaml")):
        data = yaml.safe_load(yf.read_text(encoding="utf-8"))
        if not data:
            continue
        label = '%s %s %s' % (data.get("year",""), data.get("district",""), data.get("exam_type",""))

        for q in data.get("questions", []):
            qid = q.get("id", 0)
            qtype = q.get("type", "")
            module = q.get("module", "")
            kps = q.get("knowledge_points", [])

            # 1. 解答题标记为基础/中档，但题号在26+压轴区
            if qtype == "解答":
                difficulty = q.get("difficulty", "")
                if difficulty in ("基础", "中档") and qid >= 26:
                    issues.append((label, qid, qtype, "压轴题号但标注为基础/中档",
                                   f"difficulty={difficulty}"))

            # 2. 统计/概率模块出现在解答题26+
            if qtype == "解答" and qid >= 26:
                for kp in kps:
                    if "统计" in kp or "概率" in kp:
                        issues.append((label, qid, qtype, "统计/概率出现在压轴解答题",
                                       f"kp={kp} id={qid}"))

            # 3. 解答题但模块为空
            if qtype == "解答" and not module:
                issues.append((label, qid, qtype, "解答题module字段为空", ""))

    print(f"\n{'='*60}")
    print(f"Layer B 模块-题型一致性 — math")
    print(f"{'='*60}")

    by_type = defaultdict(list)
    for item in issues:
        by_type[item[3]].append(item)

    print(f"  共发现 {len(issues)} 处一致性问题")
    for ptype, items in sorted(by_type.items(), key=lambda x: -len(x[1])):
        print(f"\n  [{ptype}] ({len(items)}处)")
        for label, qid, qtype, desc, detail in items[:5]:
            print(f"    {label} 第{qid}题: {desc} {detail}")
        if len(items) > 5:
            print(f"    ... 还有 {len(items)-5} 处")

    if not issues:
        print(f"  ✅ 无一致性问题")


# ============================================================
# Layer C: 答案提取与比对
# ============================================================

def extract_answers_from_docx(docx_path):
    """
    从解析版docx提取答案，支持多种格式。
    返回: dict[question_id] -> answer_letter
    """
    try:
        import docx
    except ImportError:
        return {}

    doc = docx.Document(docx_path)
    answers = {}
    current_q_id = None
    full_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())

    # 策略1: 表格中的答案列
    for table in doc.tables:
        if len(table.columns) < 2:
            continue
        # 找包含"答案"的列
        header_row = table.rows[0]
        ans_col_idx = None
        qid_col_idx = None
        for ci, cell in enumerate(header_row.cells):
            cell_text = cell.text.strip()
            if "答案" in cell_text:
                ans_col_idx = ci
            elif any(f"第{i}题" in cell_text or f"{i}." in cell_text for i in range(1, 30)):
                qid_col_idx = ci

        if ans_col_idx is not None:
            for row in table.rows[1:]:
                cells = row.cells
                if qid_col_idx is not None and ans_col_idx < len(cells):
                    qid_text = cells[qid_col_idx].text.strip()
                    ans_text = cells[ans_col_idx].text.strip()
                    m = re.search(r"(\d+)", qid_text)
                    if m and ans_text:
                        qid = int(m.group(1))
                        ans = ans_text.strip().upper()
                        if ans in ("A","B","C","D"):
                            answers[qid] = ans
            if answers:
                return answers  # 表格提取成功则返回

    # 策略2: 遍历段落，找"答案:"模式
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 识别题号
        m = re.match(r"^(\d+)[．.、]", text)
        if m:
            current_q_id = int(m.group(1))

        # 识别答案
        # 格式: "答案：A"
        ans_m = re.search(r"^答案[：:]\s*([A-Da-d])", text)
        if ans_m and current_q_id:
            answers[current_q_id] = ans_m.group(1).upper()
            current_q_id = None
            continue

        # 格式: "1. A" (单独一行)
        m2 = re.match(r"^\d+[．.]\s*([A-D])\s*$", text)
        if m2 and current_q_id:
            # 确认是在答案区域（前面有"答案"关键词）
            answers[current_q_id] = m2.group(1).upper()

        # 格式: 找"故选"类文字
        jieda_ans = re.search(r"(?:故选|故选择)[^A-Za-z]*([A-D])", text)
        if jieda_ans and current_q_id:
            answers[current_q_id] = jieda_ans.group(1).upper()

    return answers


def find_jiexi_path(orig_info):
    """找到对应原卷的解析版docx路径"""
    jiexi_path = orig_info.get("jiexi_path")
    if jiexi_path and Path(jiexi_path).exists():
        return jiexi_path

    # 尝试在同目录找
    folder = Path(orig_info.get("folder", ""))
    if folder.exists():
        candidates = list(folder.glob("*解析版*.docx"))
        if candidates:
            return str(candidates[0])
    return None


def build_orig_index(subject):
    """构建原卷路径索引"""
    orig_subfolder, _ = SUBJECTS[subject]
    orig_path = ORIG_DIR / orig_subfolder
    index = {}

    for top in sorted(orig_path.iterdir()):
        if not top.is_dir():
            continue
        year_match = re.match(r"(\d{4})年北京中考", top.name)
        if not year_match:
            continue
        year = year_match.group(1)

        for sub in sorted(top.iterdir()):
            if not sub.is_dir():
                continue
            fn = sub.name.replace("精品解析：", "")
            etk = None
            for dt, dk in EXAM_TYPE_MAP.items():
                if dt in fn:
                    etk = dk
                    break
            if not etk:
                continue
            dkk = None
            for dc, dk in DISTRICT_MAP.items():
                if dc in fn:
                    dkk = dk
                    break
            if not dkk:
                continue

            key = f"{year}-{dkk}-{etk}"
            jiexi_files = list(sub.glob("*解析版*.docx"))
            jiexi_path = str(jiexi_files[0]) if jiexi_files else None
            index[key] = {"folder": str(sub), "jiexi_path": jiexi_path}
    return index


def layer_c_answer_check(subject):
    """从解析版docx提取答案，与YAML比对"""
    orig_subfolder, yaml_subfolder = SUBJECTS[subject]
    yaml_path = KB_DIR / yaml_subfolder
    orig_index = build_orig_index(subject)

    yaml_files = list(yaml_path.glob("*.yaml"))
    if not yaml_files:
        print(f"[Layer C] 无YAML文件: {yaml_path}")
        return

    print(f"\n{'='*60}")
    print(f"Layer C 答案提取与比对 — {subject}")
    print(f"{'='*60}")

    total_compared = 0
    total_wrong = 0
    no_jiexi = []
    extract_fail = []
    details = []

    for yf in sorted(yaml_files):
        data = yaml.safe_load(yf.read_text(encoding="utf-8"))
        if not data:
            continue

        # 重建key
        dc = data.get("district","").replace("区","").replace("县","")
        dk = DISTRICT_MAP.get(dc, dc)
        ec = data.get("exam_type","")
        ek = EXAM_TYPE_MAP.get(ec, ec)
        key = "%s-%s-%s" % (data.get("year",""), dk, ek)

        if key not in orig_index:
            continue

        jiexi_path = orig_index[key].get("jiexi_path")
        if not jiexi_path:
            no_jiexi.append(key)
            continue

        docx_ans = extract_answers_from_docx(jiexi_path)
        if not docx_ans:
            extract_fail.append(key)
            continue

        # 比对选择/填空
        wrong = []
        for q in data.get("questions", []):
            if q.get("type") not in ("选择", "填空", "单选", "多选"):
                continue
            qid = q.get("id", 0)
            yaml_ans = q.get("answer", "").strip().upper()
            docx_ans_val = docx_ans.get(qid, "").strip().upper()
            if not docx_ans_val or not yaml_ans:
                continue
            total_compared += 1
            if yaml_ans != docx_ans_val:
                wrong.append((qid, yaml_ans, docx_ans_val))
                total_wrong += 1

        label = "%s %s %s" % (data.get("year",""), data.get("district",""), data.get("exam_type",""))
        if wrong:
            details.append((label, wrong))

    print(f"  成功提取答案的卷子: {total_compared > 0 and len(details) > 0 or '无'}")
    print(f"  共比对 {total_compared} 道选择/填空题，错误 {total_wrong} 道")
    print(f"  无法比对的原因:")
    print(f"    无解析版: {len(no_jiexi)} 份")
    print(f"    提取失败: {len(extract_fail)} 份")

    if details:
        print(f"\n  答案错误详情:")
        for label, wrong in details:
            print(f"    ❌ {label} ({len(wrong)}处错误):")
            for qid, yAns, dAns in wrong:
                print(f"       第{qid}题: YAML={yAns} 原卷={dAns}")
    else:
        print(f"  ✅ 答案全部正确（或无有效对比数据）")

    if extract_fail:
        print(f"\n  提取失败的卷子（需调解析逻辑）:")
        for k in extract_fail[:5]:
            print(f"    {k}")


# ============================================================
# 主入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(description="北京中考试卷深度质量检查")
    parser.add_argument("--subject", "-s", default="math",
                        choices=list(SUBJECTS.keys()), help="科目 (默认: math)")
    parser.add_argument("--layers", "-l", nargs="+",
                        default=["A", "B", "C"],
                        choices=["A", "B", "C"],
                        help="层级 (默认: A B C)")
    args = parser.parse_args()

    print(f"\n{'#'*60}")
    print(f"#  北京中考试卷深度质量检查")
    print(f"#  科目: {args.subject}   层: {args.layers}")
    print(f"{'#'*60}")

    for layer in args.layers:
        if layer == "A":
            layer_a_text_quality(args.subject)
        elif layer == "B":
            layer_b_module_type_consistency(args.subject)
        elif layer == "C":
            layer_c_answer_check(args.subject)

    print(f"\n{'='*60}")
    print(f"检查完成")


if __name__ == "__main__":
    main()
