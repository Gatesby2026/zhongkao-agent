#!/usr/bin/env python3

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "learning situation"
OUT_DOCX = OUT_DIR / "贾小淇_2026朝阳初三一模_学情诊断与提分建议.docx"


def load_json(path):
    return json.loads((ROOT / path).read_text(encoding="utf-8"))


def score_items(subject):
    data = load_json(f"data/answer-card-poc/{subject}/answer-card-poc.json")
    return data["scores"]["items"], data["scores"]["total_score"]


def score_value(value):
    if isinstance(value, (int, float)):
        return float(value)
    return None


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_east_asia_font(run, font_name="PingFang SC"):
    run.font.name = "Arial"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def set_style_east_asia(style, font_name="PingFang SC"):
    style.font.name = "Arial"
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    for paragraph in cell.paragraphs:
        for r in paragraph.runs:
            set_east_asia_font(r)
            r.font.size = Pt(9)


def style_table(table, header_fill="EAF2F8"):
    table.style = "Table Grid"
    table.autofit = True
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True


def add_table(doc, headers, rows, header_fill="EAF2F8"):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    style_table(table, header_fill)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_east_asia_font(run)
    return p


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.08
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.bold = bold
    set_east_asia_font(r)
    r.font.size = Pt(10.5)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        set_east_asia_font(run)
        run.font.size = Pt(10)


def compute_loss(items):
    rows = []
    total_max = 0
    total_score = 0
    for item in items:
        max_score = score_value(item["max_score"])
        score = score_value(item["score"])
        if max_score is None:
            continue
        total_max += max_score
        if score is not None:
            total_score += score
            lost = max_score - score
            if lost > 0:
                rows.append({**item, "lost": lost})
    return total_max, total_score, rows


def module_score(items, groups):
    result = []
    by_item = {item["item"]: item for item in items}
    for label, keys in groups:
        max_sum = 0
        score_sum = 0
        for key in keys:
            item = by_item.get(key)
            if not item:
                continue
            m = score_value(item["max_score"])
            s = score_value(item["score"])
            if m is not None:
                max_sum += m
            if s is not None:
                score_sum += s
        if max_sum:
            result.append((label, score_sum, max_sum, max_sum - score_sum))
    return result


def answer_map(subject):
    data = load_json(f"data/answer-card-poc/combined-structured/{subject}-best-current.json")
    mapping = {}
    for item in data["answers"]:
        number = str(item.get("number", "")).strip().rstrip(".")
        subpart = str(item.get("subpart", "")).strip()
        key = number
        if subpart:
            key = number + "_" + subpart.replace("(", "").replace(")", "").replace("：", "").replace(":", "")
        mapping.setdefault(number, []).append(item)
    return mapping


def short(text, max_len=120):
    text = str(text or "").replace("\n", " / ").strip()
    return text if len(text) <= max_len else text[: max_len - 1] + "…"


def build_report():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    math_items, math_total = score_items("math")
    chinese_items, chinese_total = score_items("chinese")
    _, _, math_lost = compute_loss(math_items)
    _, _, chinese_lost = compute_loss(chinese_items)
    math_answers = answer_map("math")
    chinese_answers = answer_map("chinese")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    set_style_east_asia(styles["Normal"])
    styles["Normal"].font.size = Pt(10.5)
    set_style_east_asia(styles["Title"])
    styles["Title"].font.size = Pt(22)
    set_style_east_asia(styles["Heading 1"])
    set_style_east_asia(styles["Heading 2"])

    header = section.header.paragraphs[0]
    header.text = "2026朝阳初三一模 | 学情诊断与提分建议"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(90, 90, 90)
        set_east_asia_font(run)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("贾小淇 2026朝阳初三一模\n学情诊断与提分建议")
    run.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)
    set_east_asia_font(run)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run(f"基于试卷题型、学生答题 OCR、小分表与评分结果生成 | {date.today().isoformat()}")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(90, 90, 90)
    set_east_asia_font(r)

    add_heading(doc, "一、总览结论", 1)
    add_table(
        doc,
        ["科目", "得分", "失分", "当前判断", "优先提分方向"],
        [
            ["数学", f"{math_total}/100", f"{100 - math_total:.1f}", "基础题稳定，后段综合题拉开差距", "函数图像、几何证明、压轴新定义"],
            ["语文", f"{chinese_total}/100", f"{100 - chinese_total:.1f}", "阅读与作文决定上限，基础有可快速修复点", "现代文答题要点、文言句意、作文升格"],
        ],
    )
    add_bullets(
        doc,
        [
            "数学的基本盘很稳：选择、填空、17-18题几乎全拿，说明基础概念和常规计算不是主要矛盾。",
            "数学真正的失分集中在 25、27、28 等综合题，属于“读题建模 + 分类讨论 + 几何证明链”能力的上限问题。",
            "语文的失分更分散：字音、默写错字、文言句意、现代文概括赏析、作文表达都有损耗；其中现代文与作文是最大提分区。",
            "后续训练不宜平均用力，应采用“基础每天少量保温 + 中高分题型集中突破 + 作文固定升格”的策略。",
        ]
    )

    add_heading(doc, "二、分数结构与失分分布", 1)
    math_modules = module_score(
        math_items,
        [
            ("选择与填空（1-16）", [str(i) for i in range(1, 17)]),
            ("基础解答（17-18）", ["17", "18"]),
            ("中档代数/应用/统计（19-23）", ["19", "20_1", "20_2", "21_1", "21_2", "22", "23_1_1", "23_1_2", "23_2_1", "23_2_2"]),
            ("综合几何与函数（24-26）", ["24_1", "24_2", "25_1", "25_2", "25_3_1", "25_3_2", "26_1", "26_2"]),
            ("压轴探究（27-28）", ["27_1", "27_2", "28_1_1", "28_1_2", "28_2"]),
        ],
    )
    add_heading(doc, "数学模块表现", 2)
    add_table(doc, ["模块", "得分", "满分", "失分", "诊断"], [
        [m, f"{s:g}", f"{mx:g}", f"{lost:g}", "稳定" if lost == 0 else ("重点突破" if lost >= 5 else "可修复")] for m, s, mx, lost in math_modules
    ])

    chinese_modules = module_score(
        chinese_items,
        [
            ("基础·运用", ["1", "2", "3", "4", "5", "6", "7"]),
            ("古诗文阅读", ["8", "9", "10", "11", "12", "13", "14", "15"]),
            ("名著阅读", ["16"]),
            ("现代文/议论文阅读", ["17", "18", "19", "20", "21", "22", "23", "24", "25", "26"]),
            ("写作", ["27_1"]),
        ],
    )
    add_heading(doc, "语文模块表现", 2)
    add_table(doc, ["模块", "得分", "满分", "失分", "诊断"], [
        [m, f"{s:g}", f"{mx:g}", f"{lost:g}", "稳定" if lost == 0 else ("重点突破" if lost >= 5 else "可修复")] for m, s, mx, lost in chinese_modules
    ])

    add_heading(doc, "三、数学错题细化分析", 1)
    math_detail = [
        ["19", "3/5", "分式化简求值；应将 a-b=3 代入得到最终数值", "过程能化到 (a-b)/2，但疑似未完成代入或结论收束", "补“条件回代”习惯；每道化简求值最后必须写“由已知得……”"],
        ["22", "4/5", "方程/实际问题建模与答题规范", "学生能列式求解，但模型或单位/检验/结论表达有小漏洞", "应用题固定四行：设、列、解、答；最后写“经检验符合题意”"],
        ["23(2)(2)", "1/2", "统计推断/可能值讨论", "答案只给出一个可能值，分类边界不完整", "做统计压轴小问时先列“数据位置/排序/边界”，再写所有可能值"],
        ["24(2)", "2/3", "圆/几何计算与证明链", "大方向可做，但几何关系或比例计算存在跳步", "证明题每一步写依据；遇切线/圆先连半径、找直角、找相似"],
        ["25(1)(3)", "1/5", "函数图像信息提取、读图、变化趋势判断", "第25题是整套数学最大失分点之一，说明图像信息转代数关系不稳", "专项练“读图三问”：坐标读值、函数关系、变化趋势/最值"],
        ["26(2)", "3/4", "二次函数参数/分类讨论", "第一问能拿，第二问少一步分类或范围论证", "二次函数题先写顶点、对称轴、开口方向；含参数必须列范围"],
        ["27(2)", "1/5", "几何综合证明与数量关系", "能写出部分关系，但辅助线和证明链不完整", "训练“图形变换题四件套”：补图、找全等/相似、写对应边角、回扣结论"],
        ["28", "0/7", "新定义/旋转线段/压轴探究", "基本空白或只写零散结论，说明读定义和套定义能力不足，也可能受时间影响", "压轴新定义先抢前两问：圈定义关键词，照例题格式代入，不纠结最后一问"],
    ]
    add_table(doc, ["题号", "得分", "采分点/参考方向", "学生问题", "提分动作"], math_detail, "FDE9D9")

    add_heading(doc, "数学提分建议", 2)
    add_bullets(
        doc,
        [
            "目标一：把 19、22、23、24、26 的小失分补回，预计可提升 5-7 分。训练重点不是刷难题，而是补齐“最后一步、分类边界、证明依据”。",
            "目标二：第25题图像函数专项突破，先把 0 分题变成 2-3 分题。每天做 1 道图像读题，要求写出读图依据。",
            "目标三：第27、28题采用“保前问”策略。27题先保证补图与第一段证明；28题先保证读懂定义、完成可直接套定义的小问。",
            "考试策略：选择填空和17-18题保持速度，不能因前面满分而放慢。把节省出的时间留给25、27、28的第一二问。",
        ]
    )

    add_heading(doc, "四、语文错题细化分析", 1)
    chinese_detail = [
        ["3", "0/2", "字音辨析；“矗立”应读 chù，不能按题面错误注音误判", "易误读字、多音字敏感度不足", "建立易错字音清单，每天10个，必须出声读+组词"],
        ["6", "1/2", "句间关系与仿写/表达连贯", "能写出内容，但逻辑关系或句式仿照不完整", "仿写先标句式骨架，再替换内容；写完检查关联词"],
        ["8", "0/1", "古诗文默写准确字形", "OCR显示“化做春泥更护花”，应为“化作”", "默写按“会背不算会，写对才算会”处理，错字单独订正"],
        ["12", "2/3", "诗词炼字/表达效果/情感", "有内容分析，但语言不够精准，情感归纳略散", "使用“三步答法”：字面义 + 画面/作用 + 情感"],
        ["14", "0/2", "文言句意理解", "语境理解偏差，可能把关键词义项代错", "文言选择题先逐字翻译，再回到句意；不能只凭语感"],
        ["15", "2/3", "文言内容概括与人物/主旨理解", "能答出部分要点，但缺少完整因果链", "答案写成“原因—行为—结果/品质”三段式"],
        ["16", "4/5", "名著人物理解深化", "能结合杨志经历，但表达较乱，人物评价前后不够收束", "名著题固定写“情节+性格+变化/启示”，避免泛泛而谈"],
        ["19-20", "2/5", "现代文信息提炼与内容概括", "概括不准或漏要点，是现代文主要失分来源", "读题圈限定词，答案尽量从原文关键句压缩改写"],
        ["22", "1.5/3", "句式特点与表达效果", "只说意思，未充分分析句式/表达效果", "套用“句式特点+内容+情感/主题作用”"],
        ["23", "3/4", "人物精神品质与启示", "观点基本对，但结合具体内容不够深入", "答开放题要有文本证据，至少两处细节支撑"],
        ["24-26", "2/7", "议论文论题、关键词、论证思路", "议论文结构意识弱，关键词理解和论证链条不清", "训练找中心论点、分论点、论据、论证方法，画文章结构图"],
        ["27作文", "34/40", "题目二“动起来”；结构完整，材料具体", "有生活经历，但错别字、比喻堆叠、结尾升华略弱", "作文重写一次：压缩重复比喻，强化转折节点和结尾立意"],
    ]
    add_table(doc, ["题号", "得分", "采分点/参考方向", "学生问题", "提分动作"], chinese_detail, "E2F0D9")

    add_heading(doc, "语文提分建议", 2)
    add_bullets(
        doc,
        [
            "基础题可快速回收 4-5 分：字音、默写、文言实词必须每天短频快复习，尤其要把错字错音写进固定清单。",
            "现代文阅读是最大提分区。每次练习不要只对答案，要把答案拆成“原文依据、概括词、表达效果”三栏，训练采分点意识。",
            "议论文阅读建议单独练 5 篇：每篇只做一件事，画出中心论点、分论点、论据、论证方法、段落作用。",
            "作文当前已在二类上段，提到 36-38 分的关键是减少语言硬伤、强化主题回扣、增加关键场景细节。",
        ]
    )

    add_heading(doc, "五、两周提分行动表", 1)
    plan_rows = [
        ["第1-2天", "整理数学错题：19、22、23、25；每题写错因和正确收束", "整理语文字音/默写/文言错题清单；作文原文通读标错字", "形成两科错题清单"],
        ["第3-5天", "每天1道函数图像题 + 1道应用题；限时并写完整步骤", "每天1篇现代文小题训练，答案拆成采分点", "补回中档题过程分"],
        ["第6-8天", "几何证明专项：24、27同类题；每步写依据", "文言句意和诗词炼字专项；按模板答", "解决跳步与表达不完整"],
        ["第9-11天", "压轴新定义/探究题：只抢前两问，练读定义", "议论文结构图训练；每篇画论证链", "提升难题入手能力"],
        ["第12-13天", "整套数学限时：前18题控时，后段保证能写采分步骤", "作文重写并润色；准备2个可迁移成长素材", "稳定发挥"],
        ["第14天", "只看错题卡和公式清单，不刷新难题", "默写、字音、作文提纲快速过一遍", "轻量复盘，避免考前焦虑"],
    ]
    add_table(doc, ["时间", "数学任务", "语文任务", "目标"], plan_rows, "D9EAF7")

    add_heading(doc, "六、打印版检查清单", 1)
    add_table(
        doc,
        ["每天完成后打勾", "内容"],
        [
            ["□", "数学：1道中档题写完整过程，不能只写答案"],
            ["□", "数学：1道函数/几何综合题标出“已知—目标—关键关系”"],
            ["□", "语文：10个易错字音/字形出声读并手写"],
            ["□", "语文：1道阅读题按“依据—概括—效果”三栏复盘"],
            ["□", "作文：积累或修改1个具体场景细节"],
        ],
    )

    add_heading(doc, "数据说明", 1)
    add_para(
        doc,
        "本报告使用学生小分表作为得分依据；学生答题内容来自答题卡大模型 OCR，作文与长解答题保留复核风险。"
        "对于未完整结构化的标准答案，报告以题目能力点、采分方向和学生失分表现进行诊断，不把 OCR 推断当作绝对标准答案。"
    )

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_report()
